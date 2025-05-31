"""
Resume Builder Module
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import io

class ResumeBuilder:
    def __init__(self):
        """Initialize the Resume Builder"""
        self.templates = {
            'Modern': self._create_modern_template,
            'Professional': self._create_professional_template,
            'Minimal': self._create_minimal_template,
            'Creative': self._create_creative_template
        }

    def generate_resume(self, resume_data: Dict[str, Any]) -> bytes:
        """Generate resume document based on template and data"""
        try:
            template_name = resume_data.get('template', 'Modern')
            template_func = self.templates.get(template_name, self._create_modern_template)
            
            # Create document
            doc = template_func(resume_data)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            print(f"Error generating resume: {str(e)}")
            return None

    def _create_modern_template(self, data: Dict[str, Any]) -> Document:
        """Create modern style resume"""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Add header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name = header.add_run(data['personal_info']['full_name'].upper())
        name.bold = True
        name.font.size = Pt(16)

        # Add contact info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = [
            data['personal_info']['email'],
            data['personal_info']['phone'],
            data['personal_info']['location']
        ]
        contact.add_run(' | '.join(filter(None, contact_info)))
        doc.add_paragraph()

        # Add summary
        if data.get('summary'):
            doc.add_heading('PROFESSIONAL SUMMARY', level=1)
            doc.add_paragraph(data['summary'])

        # Add experience
        if data.get('experience'):
            doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
            for exp in data['experience']:
                p = doc.add_paragraph()
                p.add_run(f"{exp['company']} - {exp['position']}").bold = True
                p.add_run(f"\n{exp['start_date']} - {exp['end_date']}")
                doc.add_paragraph(exp['description'])

                if exp.get('responsibilities'):
                    for resp in exp['responsibilities']:
                        doc.add_paragraph(f"• {resp}", style='List Bullet')

                if exp.get('achievements'):
                    for ach in exp['achievements']:
                        doc.add_paragraph(f"• {ach}", style='List Bullet')

        # Add education
        if data.get('education'):
            doc.add_heading('EDUCATION', level=1)
            for edu in data['education']:
                p = doc.add_paragraph()
                p.add_run(f"{edu['school']} - {edu['degree']}").bold = True
                p.add_run(f"\n{edu['field']} | {edu['graduation_date']}")
                if edu.get('gpa'):
                    p.add_run(f" | GPA: {edu['gpa']}")

        # Add skills
        if data.get('skills'):
            doc.add_heading('SKILLS', level=1)
            skills = data['skills']
            for category, skill_list in skills.items():
                if skill_list:
                    p = doc.add_paragraph()
                    p.add_run(f"{category.title()}: ").bold = True
                    p.add_run(', '.join(skill_list))

        return doc

    def _create_professional_template(self, data: Dict[str, Any]) -> Document:
        """Create professional style resume"""
        # Similar to modern template but with different styling
        return self._create_modern_template(data)

    def _create_minimal_template(self, data: Dict[str, Any]) -> Document:
        """Create minimal style resume"""
        # Similar to modern template but with minimal styling
        return self._create_modern_template(data)

    def _create_creative_template(self, data: Dict[str, Any]) -> Document:
        """Create creative style resume"""
        # Similar to modern template but with creative styling
        return self._create_modern_template(data)